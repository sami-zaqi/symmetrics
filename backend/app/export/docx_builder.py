import base64
import io

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from app.core.schemas import TestResult

DISCLAIMER = (
    "Catatan: Narasi ini adalah draf bantuan interpretasi otomatis dan wajib ditinjau serta "
    "direvisi oleh mahasiswa bersama dosen pembimbing sebelum digunakan sebagai bagian final skripsi."
)


def _set_cell_border(cell, **edges) -> None:
    """edges: e.g. top={'sz': 6}, bottom={'sz': 6}. A value of None removes the border."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge, spec in edges.items():
        el = tcBorders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tcBorders.append(el)
        if spec is None:
            el.set(qn("w:val"), "nil")
        else:
            el.set(qn("w:val"), spec.get("val", "single"))
            el.set(qn("w:sz"), str(spec.get("sz", 6)))
            el.set(qn("w:color"), spec.get("color", "000000"))


def _apply_apa_borders(table) -> None:
    """APA 7th edition table style: no vertical lines anywhere, only thin
    horizontal lines above the table, below the header row, and at the
    bottom of the table. No cell shading."""
    table.style = None  # strips any inherited grid/shading style

    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tblPr.append(borders)

    for cell in table.rows[0].cells:
        _set_cell_border(cell, top={"sz": 8}, bottom={"sz": 6})
    if len(table.rows) > 1:
        for cell in table.rows[-1].cells:
            _set_cell_border(cell, bottom={"sz": 8})


def _format_value(v) -> str:
    if v is None:
        return ""
    if isinstance(v, dict):
        if "name" in v and "value" in v:
            return f"{v['name']} = {_format_value(v['value'])}"
        return ", ".join(f"{k} = {_format_value(val)}" for k, val in v.items())
    if isinstance(v, (list, tuple)):
        if len(v) == 2 and all(isinstance(x, (int, float)) for x in v):
            return f"{_format_value(v[0])} - {_format_value(v[1])}"
        return ", ".join(_format_value(x) for x in v)
    if isinstance(v, float):
        return f"{round(v, 3):g}"
    return str(v)


def _add_dict_table(doc: Document, rows: list[dict]) -> None:
    if not rows:
        return
    keys = list(rows[0].keys())
    table = doc.add_table(rows=1, cols=len(keys))
    hdr_cells = table.rows[0].cells
    for i, k in enumerate(keys):
        hdr_cells[i].text = ""
        run = hdr_cells[i].paragraphs[0].add_run(str(k))
        run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, k in enumerate(keys):
            cells[i].text = _format_value(row.get(k))
    _apply_apa_borders(table)


def build_bab_iv_docx(result: TestResult, narrative_text: str | None) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    doc.add_heading(result.test_name_id, level=1)

    doc.add_heading("Statistik Deskriptif", level=2)
    _add_dict_table(doc, result.descriptives)

    if result.assumptions and result.assumptions.checked:
        doc.add_heading("Uji Asumsi", level=2)
        outcome_rows = [
            {
                "Uji": o.name,
                "Statistik": o.statistic,
                "p-value": o.p_value,
                "Kesimpulan": o.detail,
            }
            for o in result.assumptions.outcomes
        ]
        _add_dict_table(doc, outcome_rows)

    doc.add_heading("Hasil Uji", level=2)
    stat_rows = [{"Parameter": k, "Nilai": v} for k, v in result.test_statistics.items()]
    _add_dict_table(doc, stat_rows)

    if result.charts:
        doc.add_heading("Visualisasi", level=2)
        for chart in result.charts:
            image_bytes = base64.b64decode(chart.image_base64)
            doc.add_picture(io.BytesIO(image_bytes), width=Inches(5))

    doc.add_heading("Interpretasi", level=2)
    text = narrative_text or "(Narasi belum dibuat.)"
    if DISCLAIMER not in text:
        text = text.rstrip() + "\n\n" + DISCLAIMER
    for paragraph in text.split("\n\n"):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())

    doc.add_heading("Catatan Metodologis", level=2)
    audit_lines = [
        f"Metode: {result.method_used}",
        f"Alasan penyesuaian metode: {result.fallback_reason}" if result.fallback_reason else None,
        f"Mesin statistik: {result.engine_version}",
        f"Waktu pembuatan: {result.generated_at.isoformat()}",
    ]
    for line in audit_lines:
        if line:
            doc.add_paragraph(line, style="Intense Quote")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
